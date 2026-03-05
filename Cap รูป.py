import os
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

image_folder = r"C:\Reports" #อันนี้ใส่ชื่อโฟลเดอร์ได้เลยนะ
output_docx = r"C:\Report\Report.docx" #อันนี้ใส่ชื่อโฟลเดอร์ได้เลย
output_pdf = r"C:\Report\Report.pdf" #อันนี้ใส่ชื่อโฟลเดอร์ได้เลย

try:
    doc = Document()
    doc.add_heading('Report (แบบ PDF)' , 0)

    images = [f for f in os.listdir(image_folder) if f.endswith(('.png' , '.jpg' , '.jpeg'))]
    images.sort()

    for image_name in images:
        image_path = os.path.join(image_folder , image_name)
        doc.add_paragraph(f"Cap รูปภาพ จากหน้าเว็บ: {image_name}")
        doc.add_picture(image_path , width = Inches(6.0))
        doc.add_page_break()

        doc.save(output_docx)
        print("สร้างไฟล์ Word สำเร็จ")
        convert(output_docx , output_pdf)
        print(f"แปลงเป็น PDF สำเร็จแล้วที่: {output_pdf}")
        os.remove(output_docx)

except Exception as e:
    print(f"Error: {str(e)}")